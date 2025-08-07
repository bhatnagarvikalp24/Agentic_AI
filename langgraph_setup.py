"""
LangGraph setup for the Agentic AI Assistant.
Contains the RouterNode class and all node functions for the workflow.
"""

import json
import re
import pandas as pd
import os
from typing import Optional, List, Dict, Any
from langchain_core.runnables import Runnable
from langgraph.graph import StateGraph, END
from serpapi import GoogleSearch
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import OpenAIEmbeddings

# Import from local modules
from models import (
    GraphState, STATE_KEYS_SET_AT_ENTRY, MONEY_KEYWORDS, 
    INSURANCE_KEYWORDS, ROUTE_TYPES, CHART_TYPES
)
from config import (
    vn_model, client, embedding, call_llm, prune_state,
    SEARCH_DOMAIN_FILTER, FAISS_INDEX_PATH, SERPAPI_API_KEY
)
from prompts import (
    get_router_prompt, get_chart_suggestion_prompt, get_serp_general_summary_prompt,
    get_comparison_summary_prompt, get_faiss_summary_prompt, get_document_table_prompt
)
from database import get_schema_description, get_database_path, DATABASE_DOCUMENTATION


class RouterNode(Runnable):
    """
    Router node that determines the appropriate workflow path based on user input.
    """
    
    def invoke(self, state: GraphState, config=None) -> GraphState:
        """
        Route the user query to the appropriate node.
        
        Args:
            state: Current graph state
            config: Optional configuration
            
        Returns:
            Updated graph state with routing information
        """
        doc_flag = "yes" if state['doc_loaded'] else "no"
        db_path = get_database_path()
        schema = get_schema_description(db_path)

        # Load question-SQL pairs for examples
        qs_examples = self._load_qs_pairs()

        router_prompt = get_router_prompt(
            user_prompt=state['user_prompt'],
            doc_flag=doc_flag,
            schema=schema,
            qs_examples=qs_examples,
            documentation=DATABASE_DOCUMENTATION
        )

        try:
            response = call_llm(router_prompt)
            
            match = re.search(r'{.*}', response, re.DOTALL)
            if match:
                parsed = json.loads(match.group())
                chart_info = parsed.get("chart_info")
            else:
                print("LLM did not return valid JSON. Routing to 'search' as fallback.")
                parsed = {"route": "search"}

        except Exception as e:
            print(f"[RouterNode] LLM call failed: {e}")
            parsed = {"route": "search"}

        # Enforce safety: remove vanna_prompt if not 'document' or 'comp'
        if parsed.get("route") not in ["document", "comp", "sql", "faissdb"]:
            parsed["vanna_prompt"] = None
            parsed["fuzzy_prompt"] = None

        # Define chart_info only if needed
        chart_info = None
        
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "route": parsed.get("route"),
            "vanna_prompt": parsed.get("vanna_prompt"),
            "fuzzy_prompt": parsed.get("fuzzy_prompt"),
            "chart_info": chart_info
        }
    
    def _load_qs_pairs(self) -> str:
        """Load question-SQL pairs from file."""
        try:
            with open("./vanna_advanced_sql_pairs (1).txt", "r") as f:
                text = f.read()
            pairs = re.findall(r'question="(.*?)",\s*sql="""(.*?)"""', text, re.DOTALL)
            qs_pairs = [{"question": q.strip(), "sql": s.strip()} for q, s in pairs]
            return "\n".join(
                f"Q: {pair['question']}\nSQL: {pair['sql']}" 
                for pair in qs_pairs[:7]  # Limit to 7 to avoid token overflow
            )
        except Exception as e:
            print(f"Error loading QS pairs: {e}")
            return ""


def get_user_chart_type(prompt: str) -> Optional[str]:
    """
    Extract user-specified chart type from prompt.
    
    Args:
        prompt: User's input prompt
        
    Returns:
        Chart type if specified, None otherwise
    """
    prompt = prompt.lower()
    if "bar chart" in prompt or "bar graph" in prompt:
        return "bar"
    elif "line chart" in prompt or "line graph" in prompt:
        return "line"
    elif "pie chart" in prompt or "pie graph" in prompt:
        return "pie"
    return None


def suggest_chart(df: pd.DataFrame) -> Optional[dict]:
    """
    Suggest appropriate chart type based on data.
    
    Args:
        df: DataFrame to analyze
        
    Returns:
        Chart configuration dictionary
    """
    sample_data = df.head(5).to_dict(orient="list")
    prompt = get_chart_suggestion_prompt(json.dumps(sample_data, indent=2))

    reply = call_llm(prompt)
    match = re.search(r'{.*}', reply, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except:
            return None
    return None


def plot_chart(df: pd.DataFrame, chart_info: dict):
    """
    Plot chart based on configuration.
    
    Args:
        df: DataFrame to plot
        chart_info: Chart configuration
    """
    chart_type = chart_info.get("type", "bar")
    x = chart_info.get("x")
    y = chart_info.get("y")

    if isinstance(y, str):
        y = [y]  # Make it a list

    df_columns = list(df.columns)
    def match_col(col_name):
        for c in df_columns:
            if col_name.lower().replace(" ", "") in c.lower().replace(" ", ""):
                return c
        return None

    x_col = match_col(x)
    y_cols = [match_col(col) for col in y if match_col(col)]

    if not x_col or not y_cols:
        print(f"Invalid chart columns: {x}, {y}")
        return

    print(f"{chart_type.capitalize()} Chart: {', '.join(y)} vs {x}")

    if chart_type == "bar":
        # Implementation for bar chart
        pass
    elif chart_type == "line":
        # Implementation for line chart
        pass
    elif chart_type == "pie" and len(y_cols) == 1:
        # Implementation for pie chart
        pass
    else:
        print("Pie chart supports only one y column.")


def vanna_node(state: GraphState) -> GraphState:
    """
    Node for executing Vanna SQL queries.
    
    Args:
        state: Current graph state
        
    Returns:
        Updated state with SQL results
    """
    prompt = state["vanna_prompt"]

    sql_query = vn_model.generate_sql(prompt)

    try:
        result = vn_model.run_sql(sql_query)
        if isinstance(result, pd.DataFrame):
            parsed_result = result
        elif isinstance(result, list):
            parsed_result = pd.DataFrame(result)
        else:
            parsed_result = pd.DataFrame([{"Result": str(result)}])
    except Exception as e:
        parsed_result = pd.DataFrame([{"Error": f"SQL Execution failed: {e}"}])

    return {
        **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
        "sql_result": parsed_result,
        "sql_query": sql_query
    }


def enhance_query(prompt: str) -> str:
    """
    Enhance search query with insurance context and domain filters.
    
    Args:
        prompt: Original user prompt
        
    Returns:
        Enhanced search query
    """
    if any(keyword in prompt.lower() for keyword in INSURANCE_KEYWORDS):
        base_query = prompt
    else:
        base_query = f"In the insurance industry, {prompt}"
    
    return f"{base_query} {SEARCH_DOMAIN_FILTER}"


def serp_node(state: GraphState) -> GraphState:
    """
    Node for performing web searches using SerpAPI.
    
    Args:
        state: Current graph state
        
    Returns:
        Updated state with search results
    """
    query = enhance_query(state["user_prompt"])

    search = GoogleSearch({
        "q": query,
        "api_key": SERPAPI_API_KEY,
        "num": 5
    })
    results = search.get_dict()

    links = []
    summaries = []

    if "organic_results" in results:
        for r in results["organic_results"][:5]:
            link = r.get("link")
            title = r.get("title", "Untitled").strip('"')
            snippet = r.get("snippet", "No summary available.").strip('"')
            if link:
                links.append(f"[{title}]({link})")
                summaries.append(snippet)

    if not links:
        links = ["No insurance-related results found or API limit reached."]
        summaries = [""]

    # Add LLM-generated general summary with numeric insights
    combined_text = "\n".join(summaries)

    # Build conditional prompt for COMP vs SERP node
    if "sql_query" in state and state["sql_query"]:
        sql_snippet = f"\n🧾 Internal SQL Query:\n{state['sql_query']}"
    else:
        sql_snippet = ""

    if "sql_result" in state and isinstance(state["sql_result"], pd.DataFrame):
        sql_snippet += f"\n\n📊 Top 5 rows of SQL Output:\n{state['sql_result'].head(5).to_markdown(index=False)}"

    general_summary_prompt = get_serp_general_summary_prompt(
        combined_text=combined_text,
        user_prompt=state['user_prompt'],
        sql_snippet=sql_snippet
    )

    general_summary = call_llm(general_summary_prompt)

    return {
        **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
        "web_links": list(zip(links, summaries)),
        "general_summary": general_summary
    }


def comp_node(state: GraphState) -> GraphState:
    """
    Node for comparison analysis between internal and external data.
    
    Args:
        state: Current graph state
        
    Returns:
        Updated state with comparison results
    """
    # Step 1: Run Vanna SQL
    vanna_prompt = state.get("vanna_prompt") or state["user_prompt"]
    sql_query = vn_model.generate_sql(vanna_prompt)

    try:
        result = vn_model.run_sql(sql_query)
        if isinstance(result, pd.DataFrame):
            parsed_result = result
        elif isinstance(result, list):
            parsed_result = pd.DataFrame(result)
        else:
            parsed_result = pd.DataFrame([{"Result": str(result)}])
    except Exception as e:
        parsed_result = pd.DataFrame([{"Error": f"SQL Execution failed: {e}"}])
    
    sql_df = parsed_result
    
    # Step 2: Run Serp Search
    serp_result = serp_node({**state, "sql_query": sql_query, "sql_result": sql_df})
    web_links = serp_result.get("web_links")
    external_summary = serp_result.get("general_summary", "")

    # Step 3: Generate comparison summary using LLM
    summary_prompt = get_comparison_summary_prompt(
        sql_df=sql_df.head(5).to_markdown(index=False) if isinstance(sql_df, pd.DataFrame) else str(sql_df),
        web_links="\n".join([f"- {title}: {summary[:200]}..." for title, summary in web_links]),
        external_summary=external_summary
    )
    
    comparison_summary = call_llm(summary_prompt)

    return {
        **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
        "sql_result": sql_df,
        "sql_query": sql_query,
        "web_links": web_links,
        "general_summary": serp_result.get("general_summary", ""),
        "comparison_summary": comparison_summary
    }


def faissdb_node(state: GraphState) -> GraphState:
    """
    Node for querying internal knowledge base using FAISS.
    
    Args:
        state: Current graph state
        
    Returns:
        Updated state with FAISS results
    """
    faiss = FAISS.load_local(
        folder_path=FAISS_INDEX_PATH,
        embeddings=OpenAIEmbeddings(),
        allow_dangerous_deserialization=True
    )
    docs = faiss.similarity_search(state["user_prompt"], k=3)

    top_docs = docs[:3]
    content_snippets = "\n\n---\n\n".join(d.page_content[:500] for d in top_docs)

    summary_prompt = get_faiss_summary_prompt(
        user_prompt=state['user_prompt'],
        content_snippets=content_snippets
    )
    summary = call_llm(summary_prompt)

    # Extract faiss_sources with source path
    faiss_sources = []
    all_images = []

    for doc in top_docs:
        doc_name = doc.metadata.get("source_doc", "Doc")
        snippet = doc.page_content[:300]
        path = doc.metadata.get("file_path")
        faiss_sources.append((doc_name, snippet, path))

        # Load associated images
        image_meta_path = os.path.join("extracted_images", "extracted_image_metadata.json")
        if os.path.exists(image_meta_path):
            with open(image_meta_path, 'r') as f:
                all_metadata = json.load(f)
            related_images = [
                meta for meta in all_metadata
                if meta["original_doc"] == doc_name
            ]
            all_images.extend(related_images)

    return {
        **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
        "faiss_summary": summary,
        "faiss_sources": faiss_sources,
        "faiss_images": all_images
    }


def document_node(state: GraphState) -> GraphState:
    """
    Node for processing and updating documents.
    
    Args:
        state: Current graph state
        
    Returns:
        Updated state with document processing results
    """
    doc_path = state['document_path']
    if not doc_path or not state.get("vanna_prompt"):
        return {
            **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
            "updated_doc_path": None
        }

    doc = Document(doc_path)

    structure_string = ""
    header = None
    header_table_map = {}

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            header = para.text.strip()
            structure_string += f"\n# {header}"
            header_table_map[header] = []
        elif header:
            header_table_map[header].append(len(header_table_map[header]))

    for idx, table in enumerate(doc.tables):
        cols = [cell.text.strip() for cell in table.rows[0].cells]
        structure_string += f"\n- Table {idx}: {len(table.rows)} rows x {len(cols)} columns, Columns: {cols}"

    prompt = get_document_table_prompt(
        structure_string=structure_string,
        fuzzy_prompt=state['fuzzy_prompt']
    )
    llm_output = call_llm(prompt)
    json_match = re.search(r'{.*}', llm_output, re.DOTALL)
    parsed = json.loads(json_match.group()) if json_match else {"header_text": list(header_table_map)[0], "table_index_under_header": 0}

    # Generate SQL via Vanna
    try:
        sql_query = vn_model.generate_sql(state["vanna_prompt"])
        vanna_output = vn_model.run_sql(sql_query)
    except Exception as e:
        return {**state, "updated_doc_path": None, "error": f"SQL generation or execution failed: {e}"}

    # Update the correct table
    header = parsed['header_text']
    table_idx = parsed['table_index_under_header']
    matched_table_index = list(header_table_map[header])[table_idx]
    table = doc.tables[matched_table_index]

    # Fill table with SQL output
    if isinstance(vanna_output, pd.DataFrame):
        for i, row in enumerate(vanna_output.itertuples(index=False), start=1):
            for j, value in enumerate(row):
                if i < len(table.rows) and j < len(table.columns):
                    table.cell(i, j).text = str(value)

    updated_path = "updated_doc.docx"
    doc.save(updated_path)

    return {
        **prune_state(state, STATE_KEYS_SET_AT_ENTRY),
        "updated_doc_path": updated_path,
        "header_updated": header,
        "table_index_updated": matched_table_index
    }


def build_langgraph() -> StateGraph:
    """
    Build and configure the LangGraph workflow.
    
    Returns:
        Compiled LangGraph instance
    """
    # Create graph builder
    graph_builder = StateGraph(GraphState)
    
    # Add nodes
    graph_builder.add_node("router", RouterNode())
    graph_builder.add_node("vanna_sql", vanna_node)
    graph_builder.add_node("serp_search", serp_node)
    graph_builder.add_node("doc_update", document_node)
    graph_builder.add_node("comp", comp_node)
    graph_builder.add_node("faissdb", faissdb_node)

    # Define routing logic
    def router_logic(state: GraphState):
        if state['route'] == 'sql': 
            return "vanna_sql"
        elif state['route'] == 'search': 
            return "serp_search"
        elif state['route'] == 'document': 
            return "doc_update"
        elif state['route'] == 'comp': 
            return "comp"
        elif state['route'] == 'faissdb': 
            return "faissdb"
        else: 
            return END    

    # Set entry point
    graph_builder.set_entry_point("router")

    # Add conditional edges
    graph_builder.add_conditional_edges("router", router_logic)

    # Add edges to end
    graph_builder.add_edge("vanna_sql", END)
    graph_builder.add_edge("serp_search", END)
    graph_builder.add_edge("doc_update", END)
    graph_builder.add_edge("comp", END)
    graph_builder.add_edge("faissdb", END)

    # Compile and return
    return graph_builder.compile() 